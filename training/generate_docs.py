"""Generate the BillSense counterfeit-detection documentation as .docx AND .pdf
into Resources/Documents/. Content mirrors docs/HOW_COUNTERFEIT_DETECTION_WORKS.md.
Uses python-docx (Word) + reportlab (PDF) — no pandoc/LibreOffice needed."""
import os

OUT = r"D:\Github\Billsense-Project\Resources\Documents"
BASE = "BillSense_How_Scanning_and_Counterfeit_Detection_Works"
os.makedirs(OUT, exist_ok=True)

# ---- content model: list of blocks -------------------------------------------------
# ('h1',txt) ('h2',txt) ('p',txt) ('b',[items]) ('num',[items]) ('table',[rows]) ('note',txt)
DOC = [
    ('title', "BillSense — How Scanning & Counterfeit Detection Work"),
    ('sub',   "How the models work • what the system measures • old-vs-new model metrics"),
    ('sub',   "Updated 2026-06-18  ·  Server API v17.13  ·  Denomination v2 + securitycf v1"),

    ('h1', "0. Update — Real Counterfeit Detection is Now Live"),
    ('p', "As of server API v17.13, BillSense has a dedicated counterfeit model, securitycf.pt "
          "(16 classes: 8 genuine security features + 8 FALSE markers), trained on labelled real AND "
          "fake banknote features. The scan now:"),
    ('b', [
        "Detects watermark, see-through, shadow thread, security thread, concealed value, OVI and EVP "
        "(the security models previously lacked the labelled data for the hard ones).",
        "Flags forgeries directly: any FALSE marker (false watermark, false security thread, false OVI, "
        "false bill, etc.) makes the verdict COUNTERFEIT and names the marker.",
    ]),
    ('p', "So the system below — previously a \"genuine-feature verifier\" — now also actively catches "
          "fakes. The sections that follow describe the full pipeline; the securitycf additions are noted "
          "in the relevant places."),

    ('h1', "1. The Honest Headline"),
    ('p', "BillSense does NOT have a single real-vs-fake AI model. There is no model you can hand a "
          "photo and get back \"genuine\" or \"counterfeit.\" The verdict is COMPUTED from two kinds of evidence:"),
    ('num', [
        "A specific learned forgery marker — the EVP model (evp.pt) was trained with explicit "
        "\"false 1k / 500 enhanced value panel\" classes. If detected, the note is flagged COUNTERFEIT. "
        "This is the only place a model directly says \"fake\", and it only covers the enhanced value "
        "panel on PHP 500 / PHP 1000 notes.",
        "Genuine-feature verification — the other models find security features (security thread, serial "
        "number, concealed value, watermark, OVI, etc.). More features verified = higher authenticity "
        "score. Few features = lower confidence, NOT an accusation.",
    ]),
    ('p', "In short, the system today is a \"genuine-feature verifier + one specific forgery-marker "
          "detector\" — not a general counterfeit classifier."),
    ('b', [
        "It reliably identifies the denomination and verifies genuine security features.",
        "It catches a fake that carries a false enhanced value panel.",
        "It will NOT reliably catch a high-quality fake that has decent features and no false-EVP marker "
        "— such a note reads as LIKELY GENUINE.",
        "A single front-lit phone photo physically cannot see the watermark/see-through (needs backlight) "
        "or OVI/OVD (needs tilt); their absence is treated as \"not verified yet\", never as proof of forgery.",
    ]),
    ('note', "Deliberate trade-off: condemning a note on weak evidence would false-accuse real money. "
             "BillSense errs toward \"verify what we can; only condemn on a positive forgery signal.\""),

    ('h1', "2. The Six Models"),
    ('p', "All six are YOLOv8 detectors. Each finds objects; none outputs a genuine/fake probability "
          "except the EVP model's \"false ...\" classes."),
    ('table', [
        ["Model file", "Detects", "Role"],
        ["denomination2.pt *", "20, 50, 100, 200, 500, 1000", "Identifies the note value. Gates everything. (retrained)"],
        ["security_best.pt", "security thread, serial number, concealed value, watermark, see-through, value", "Finds genuine security features (OBB)"],
        ["counterfeit_best.pt", "UV-thread, concealed-value, security-thread, serial-number, symbol-of-nature, value", "Despite the name, a FEATURE detector — not a fake classifier"],
        ["ovi.pt", "optically variable ink", "Colour-shifting ink (high-denom)"],
        ["ovd.pt", "ovd", "Optically variable device (high-denom)"],
        ["evp.pt (!)", "1k/500 EVP, FALSE 1k/500 EVP", "The only forgery detector — the \"false\" classes are the counterfeit signal"],
    ]),
    ('p', "* Only denomination2.pt was retrained in this cycle. (!) evp.pt is the lone source of a "
          "positive \"counterfeit\" signal."),

    ('h1', "3. The Scan Pipeline (single-photo standard scan)"),
    ('num', [
        "Decode the photo.",
        "Denomination detection (denomination2.pt, confidence >= 0.25). If UNKNOWN -> verdict = UNKNOWN "
        "(\"re-scan\") and the pipeline stops.",
        "Security-feature detection (parallel): security_best, counterfeit_best, ovi, ovd, evp, plus "
        "classic watermark analysis. Coverage = detected features / expected (6 for low-denom, 9 for "
        "high-denom PHP 500/1000). The false-EVP marker is captured here.",
        "Capture-quality analysis (blur / brightness / contrast).",
        "Feature-geometry measurement (each feature's placement vs empirical reference medians).",
        "evaluate_counterfeit() -> status + 0-100 authenticity score, plus an annotated image stored to Firebase.",
    ]),
    ('p', "Multi-Scan mode additionally measures colour-shift across tilt angles to verify OVI/OVD — the "
          "features a single photo cannot capture — and can boost the verdict."),

    ('h1', "4. How the Verdict Is Decided"),
    ('table', [
        ["Condition", "Result"],
        ["Denomination not identified", "UNKNOWN — \"re-scan the full note\""],
        ["False enhanced value panel detected", "COUNTERFEIT (score capped <= 15)"],
        ["EVP present on a low-denom note", "COUNTERFEIT"],
        ["Capture quality < 30 (blurry/dark)", "NEEDS_RESCAN"],
        ["otherwise", "compute score, then tier it (below)"],
    ]),
    ('h2', "The authenticity score (0-100)"),
    ('p', "base = 0.60 x feature-coverage + 0.40 x detection-confidence;  "
          "base += 0.15 x geometry-placement (BONUS only, never lowers a genuine note);  "
          "score = 100 x base x quality-factor   (quality-factor = clamp(quality/70, 0.4, 1.0))."),
    ('table', [
        ["Score", "Status"],
        [">= 75", "GENUINE (high confidence)"],
        [">= 50", "GENUINE (medium)"],
        ["< 50", "LIKELY GENUINE (few features visible — use Multi-Scan)"],
    ]),
    ('note', "Guiding rule (v17.1): missing features != counterfeit. COUNTERFEIT is reserved for a "
             "positive forgery signal (false EVP). Everything else is scored by how much genuine evidence "
             "was gathered — which is why a clean photo of a real note with few visible features reads "
             "LIKELY GENUINE, not COUNTERFEIT."),

    ('h1', "5. Model Score Metrics — Old vs New (Denomination Model)"),
    ('p', "The 2026-06-18 retrain added two new full-resolution datasets (PH Banknote.coco + 2025 "
          "\"Peso bill Detector\" captures) and rebalanced the train/val split. Measured with Ultralytics "
          "on the held-out validation set (482 images neither model trained on):"),
    ('table', [
        ["Metric", "OLD", "NEW", "Change"],
        ["mAP@50", "0.726", "0.928", "+0.20"],
        ["mAP@50-95", "0.558", "0.818", "+0.26"],
        ["Precision", "0.763", "0.924", "+0.16"],
        ["Recall", "0.646", "0.867", "+0.22"],
    ]),
    ('h2', "Per-denomination (mAP@50 / recall)"),
    ('table', [
        ["Denom", "OLD mAP50", "NEW mAP50", "OLD recall", "NEW recall"],
        ["PHP 20", "0.671", "0.993", "0.375", "0.987"],
        ["PHP 50", "0.571", "0.949", "0.588", "0.917"],
        ["PHP 100", "0.880", "0.920", "0.839", "0.903"],
        ["PHP 200", "0.894", "0.938", "0.776", "0.821"],
        ["PHP 500", "0.695", "0.912", "0.619", "0.785"],
        ["PHP 1000", "0.646", "0.858", "0.681", "0.792"],
    ]),
    ('note', "These metrics are for the DENOMINATION detector only — the model retrained this cycle. The "
             "feature/forgery models (security_best, counterfeit_best, ovi, ovd, evp) were not retrained. "
             "Better denomination recall matters because an unidentified note (UNKNOWN) produces no verdict "
             "at all — so this retrain directly increases how often the counterfeit-evaluation pipeline can "
             "run (e.g. PHP 50 went from frequently-UNKNOWN to reliably identified)."),

    ('h1', "6. Making Counterfeit Detection Stronger (Future Work)"),
    ('num', [
        "Train a dedicated realfake.pt classifier on labelled genuine AND counterfeit examples, wired into "
        "evaluate_counterfeit() as an additional vote (never the sole decider, to avoid false accusations).",
        "Expand the forgery-marker classes beyond the false EVP (e.g. false security thread, mismatched "
        "serial fonts) so more fake types trip a positive signal.",
        "Use Multi-Scan for OVI/OVD/watermark — the tilt/backlit features are the hardest to fake; "
        "verifying them across angles is the strongest available genuine signal.",
        "UV hardware — true UV-reactive features need a UV light source the phone camera lacks.",
    ]),

    ('h1', "7. TL;DR"),
    ('b', [
        "BillSense identifies the bill and verifies genuine security features, producing a 0-100 "
        "authenticity score and a status: GENUINE / LIKELY GENUINE / NEEDS_RESCAN / COUNTERFEIT / UNKNOWN.",
        "The only direct \"it's fake\" signal today is the false enhanced value panel (PHP 500/1000). "
        "Otherwise the system measures genuineness; low evidence means low confidence, not \"fake.\"",
        "The 2026-06-18 retrain lifted the denomination detector from mAP50 0.73 to 0.93, letting the rest "
        "of the pipeline run on notes (like PHP 50) it used to fail to identify.",
        "A general counterfeit classifier (realfake.pt) is future work, not a current feature.",
    ]),
]

# ---------------- DOCX renderer ----------------
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

d = Document()
st = d.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
NAVY = RGBColor(0x1F, 0x3A, 0x6E); BLUE = RGBColor(0x2E, 0x75, 0xB6)

def shade(cell, hexcolor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor); tcPr.append(sh)

for kind, *rest in (b if isinstance(b, tuple) else (b,) for b in DOC):
    data = rest[0]
    if kind == 'title':
        p = d.add_paragraph(); r = p.add_run(data); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
    elif kind == 'sub':
        p = d.add_paragraph(); r = p.add_run(data); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    elif kind == 'h1':
        p = d.add_paragraph(); r = p.add_run(data); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    elif kind == 'h2':
        p = d.add_paragraph(); r = p.add_run(data); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(8)
    elif kind == 'p':
        d.add_paragraph(data)
    elif kind == 'note':
        p = d.add_paragraph(); r = p.add_run("Note: " + data); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    elif kind == 'b':
        for it in data: d.add_paragraph(it, style='List Bullet')
    elif kind == 'num':
        for it in data: d.add_paragraph(it, style='List Number')
    elif kind == 'table':
        rows = data; t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = 'Table Grid'
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri].cells[ci]; cell.text = ""
                rr = cell.paragraphs[0].add_run(val); rr.font.size = Pt(9.5)
                if ri == 0:
                    rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); shade(cell, "2E75B6")
                elif ri % 2 == 0:
                    shade(cell, "EEF3FA")

docx_path = os.path.join(OUT, BASE + ".docx")
d.save(docx_path); print("DOCX ->", docx_path)

# ---------------- PDF renderer ----------------
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                                ListFlowable, ListItem)

ss = getSampleStyleSheet()
navy = colors.HexColor("#1F3A6E"); blue = colors.HexColor("#2E75B6")
H1 = ParagraphStyle('H1', parent=ss['Heading1'], textColor=navy, fontSize=15, spaceBefore=14, spaceAfter=4)
H2 = ParagraphStyle('H2', parent=ss['Heading2'], textColor=blue, fontSize=12, spaceBefore=8)
BODY = ParagraphStyle('Body', parent=ss['BodyText'], fontSize=10.5, leading=15)
NOTE = ParagraphStyle('Note', parent=BODY, fontSize=9.5, textColor=colors.HexColor("#555555"), leftIndent=10, borderColor=blue)
TITLE = ParagraphStyle('Title2', parent=ss['Title'], textColor=navy, fontSize=20)
SUB = ParagraphStyle('Sub', parent=BODY, fontSize=10, textColor=colors.HexColor("#666666"))

def esc(s): return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

flow = []
for kind, *rest in (b if isinstance(b, tuple) else (b,) for b in DOC):
    data = rest[0]
    if kind == 'title': flow += [Paragraph(esc(data), TITLE), Spacer(1,4)]
    elif kind == 'sub': flow += [Paragraph(esc(data), SUB)]
    elif kind == 'h1': flow += [Paragraph(esc(data), H1)]
    elif kind == 'h2': flow += [Paragraph(esc(data), H2)]
    elif kind == 'p': flow += [Paragraph(esc(data), BODY), Spacer(1,3)]
    elif kind == 'note': flow += [Spacer(1,2), Paragraph("<i>Note: " + esc(data) + "</i>", NOTE), Spacer(1,3)]
    elif kind in ('b','num'):
        items = [ListItem(Paragraph(esc(x), BODY), leftIndent=18) for x in data]
        flow += [ListFlowable(items, bulletType=('bullet' if kind=='b' else '1'), start=('•' if kind=='b' else '1')), Spacer(1,4)]
    elif kind == 'table':
        rows = [[Paragraph(esc(c), ParagraphStyle('cell', parent=BODY, fontSize=8.7,
                 textColor=(colors.white if ri==0 else colors.black), leading=11))
                 for c in row] for ri, row in enumerate(data)]
        tbl = Table(rows, repeatRows=1)
        sty = [('BACKGROUND',(0,0),(-1,0),blue), ('GRID',(0,0),(-1,-1),0.5,colors.HexColor("#CCCCCC")),
               ('VALIGN',(0,0),(-1,-1),'MIDDLE'), ('TOPPADDING',(0,0),(-1,-1),4),
               ('BOTTOMPADDING',(0,0),(-1,-1),4), ('LEFTPADDING',(0,0),(-1,-1),6)]
        for ri in range(1,len(rows)):
            if ri % 2 == 0: sty.append(('BACKGROUND',(0,ri),(-1,ri),colors.HexColor("#EEF3FA")))
        tbl.setStyle(TableStyle(sty)); flow += [tbl, Spacer(1,6)]

pdf_path = os.path.join(OUT, BASE + ".pdf")
SimpleDocTemplate(pdf_path, pagesize=letter, topMargin=0.8*inch, bottomMargin=0.8*inch,
                  leftMargin=0.9*inch, rightMargin=0.9*inch,
                  title="BillSense — Scanning & Counterfeit Detection").build(flow)
print("PDF  ->", pdf_path)
